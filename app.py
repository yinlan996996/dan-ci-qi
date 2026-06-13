"""
单词筛查背诵应用 - 后端
Flask + SQLite，轻量、单文件、一条命令启动
"""

import csv
import io
import json
import os
import random
import sqlite3
from datetime import datetime

from flask import Flask, g, jsonify, render_template, request, send_file

# ─── 初始化 ───────────────────────────────────────────────
app = Flask(__name__)

DATABASE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "words.db")


def get_db():
    """获取数据库连接"""
    if "db" not in g:
        g.db = sqlite3.connect(DATABASE)
        g.db.row_factory = sqlite3.Row
        g.db.execute("PRAGMA foreign_keys = ON")
    return g.db


@app.teardown_appcontext
def close_db(exception):
    db = g.pop("db", None)
    if db is not None:
        db.close()


def init_db():
    """建表"""
    db = get_db()
    db.executescript("""
        CREATE TABLE IF NOT EXISTS words (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            english TEXT NOT NULL UNIQUE,
            chinese TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS quiz_sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            total_count INTEGER NOT NULL DEFAULT 0,
            correct_count INTEGER NOT NULL DEFAULT 0,
            wrong_count INTEGER NOT NULL DEFAULT 0,
            word_order TEXT NOT NULL DEFAULT '[]',
            current_index INTEGER NOT NULL DEFAULT 0,
            status TEXT NOT NULL DEFAULT 'in_progress',
            mode TEXT NOT NULL DEFAULT 'all',
            started_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            completed_at TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS quiz_answers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id INTEGER NOT NULL,
            word_id INTEGER NOT NULL,
            is_correct BOOLEAN NOT NULL,
            user_answer TEXT NOT NULL DEFAULT '',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (session_id) REFERENCES quiz_sessions(id),
            FOREIGN KEY (word_id) REFERENCES words(id)
        );
    """)
    db.commit()


# 首次请求时初始化数据库
@app.before_request
def before_request():
    init_db()


# ─── 首页 ──────────────────────────────────────────────────
@app.route("/")
def index():
    return render_template("index.html")


# ═══════════════════════════════════════════════════════════════
#  单词管理 API
# ═══════════════════════════════════════════════════════════════

@app.route("/api/words", methods=["GET"])
def get_words():
    """获取所有单词，支持 ?status=error 筛选错题单词"""
    db = get_db()
    status = request.args.get("status", "")

    if status == "error":
        rows = db.execute("""
            SELECT w.*, COUNT(qa.id) AS error_count
            FROM words w
            INNER JOIN quiz_answers qa ON w.id = qa.word_id AND qa.is_correct = 0
            GROUP BY w.id
            ORDER BY error_count DESC
        """).fetchall()
    else:
        rows = db.execute(
            "SELECT *, 0 AS error_count FROM words ORDER BY created_at DESC"
        ).fetchall()

    return jsonify([dict(r) for r in rows])


@app.route("/api/words", methods=["POST"])
def add_word():
    """手动添加单个单词"""
    data = request.get_json(silent=True) or {}
    english = (data.get("english") or "").strip()
    chinese = (data.get("chinese") or "").strip()

    if not english or not chinese:
        return jsonify({"error": "英文和中文不能为空"}), 400

    db = get_db()
    try:
        cur = db.execute(
            "INSERT INTO words (english, chinese) VALUES (?, ?)",
            (english, chinese),
        )
        db.commit()
        return jsonify({"id": cur.lastrowid, "english": english, "chinese": chinese}), 201
    except sqlite3.IntegrityError:
        return jsonify({"error": f"单词 '{english}' 已存在"}), 409


@app.route("/api/words/<int:word_id>", methods=["DELETE"])
def delete_word(word_id):
    """删除单词"""
    db = get_db()
    db.execute("DELETE FROM words WHERE id = ?", (word_id,))
    db.commit()
    return jsonify({"ok": True})


@app.route("/api/words/all", methods=["DELETE"])
def clear_words():
    """清空所有单词（连带清除答题记录）"""
    db = get_db()
    db.execute("DELETE FROM quiz_answers")
    db.execute("DELETE FROM quiz_sessions")
    db.execute("DELETE FROM words")
    db.commit()
    return jsonify({"ok": True})


@app.route("/api/words/import", methods=["POST"])
def import_words():
    """批量导入单词 — 支持 JSON body、文件上传、文本粘贴"""
    db = get_db()
    words_to_add = []
    errors = []

    # ── 方式1：JSON body ──
    if request.is_json:
        data = request.get_json()
        if isinstance(data, list):
            words_to_add = data
        elif isinstance(data, dict) and "words" in data:
            words_to_add = data["words"]

    # ── 方式2：文件上传 ──
    elif request.files:
        file = None
        for key in request.files:
            file = request.files[key]
            break

        if file:
            filename = file.filename.lower()
            try:
                if filename.endswith(".json"):
                    content = json.loads(file.read().decode("utf-8"))
                    words_to_add = content if isinstance(content, list) else content.get("words", [])
                elif filename.endswith(".csv"):
                    text = file.read().decode("utf-8-sig")
                    reader = csv.reader(io.StringIO(text))
                    for row in reader:
                        if len(row) >= 2 and row[0].strip() and row[1].strip():
                            words_to_add.append({"english": row[0].strip(), "chinese": row[1].strip()})
                elif filename.endswith(".xlsx"):
                    import openpyxl
                    wb = openpyxl.load_workbook(io.BytesIO(file.read()))
                    ws = wb.active
                    for row in ws.iter_rows(min_row=2 if ws.cell(1, 1).value and ws.cell(1, 1).value in ("english", "英文", "单词") else 1):
                        eng = str(row[0].value).strip() if row[0].value else ""
                        chn = str(row[1].value).strip() if len(row) > 1 and row[1].value else ""
                        if eng and chn and eng != "english" and eng != "英文":
                            words_to_add.append({"english": eng, "chinese": chn})
            except Exception as e:
                return jsonify({"error": f"文件解析失败: {str(e)}"}), 400

    # ── 方式3：文本粘贴 ──
    else:
        text = (request.form.get("text") or request.get_data(as_text=True)).strip()
        if text:
            for line in text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                # 用 tab、多个空格、或 " - " 分隔
                parts = None
                if "\t" in line:
                    parts = line.split("\t", 1)
                elif "  " in line:
                    parts = line.split("  ", 1)
                elif " - " in line:
                    parts = line.split(" - ", 1)
                else:
                    parts = line.split(None, 1)  # 第一个空白分隔
                if parts and len(parts) == 2:
                    words_to_add.append({"english": parts[0].strip(), "chinese": parts[1].strip()})

    if not words_to_add:
        return jsonify({"error": "未识别到有效单词数据"}), 400

    # 逐条插入
    added_count = 0
    skipped_count = 0
    for item in words_to_add:
        eng = (item.get("english") or "").strip()
        chn = (item.get("chinese") or "").strip()
        if not eng or not chn:
            continue
        try:
            db.execute(
                "INSERT INTO words (english, chinese) VALUES (?, ?)",
                (eng, chn),
            )
            added_count += 1
        except sqlite3.IntegrityError:
            skipped_count += 1
        except Exception as e:
            errors.append(f"{eng}: {str(e)}")

    db.commit()
    return jsonify({
        "ok": True,
        "added": added_count,
        "skipped": skipped_count,
        "errors": errors,
    })


# ═══════════════════════════════════════════════════════════════
#  筛查测试 API
# ═══════════════════════════════════════════════════════════════

@app.route("/api/quiz/start", methods=["POST"])
def quiz_start():
    """开始新筛查"""
    db = get_db()
    mode = request.args.get("mode", "all")

    if mode == "errors":
        # 获取所有错过的单词 ID
        error_rows = db.execute("""
            SELECT DISTINCT w.id, w.english, w.chinese
            FROM words w
            INNER JOIN quiz_answers qa ON w.id = qa.word_id AND qa.is_correct = 0
        """).fetchall()
        word_list = [dict(r) for r in error_rows]
    else:
        word_list = [dict(r) for r in db.execute("SELECT id, english, chinese FROM words").fetchall()]

    if len(word_list) < 1:
        return jsonify({"error": "没有可筛查的单词，请先导入"}), 400
    if len(word_list) < 4:
        return jsonify({"error": f"至少需要 4 个单词才能生成四选一题目（当前只有 {len(word_list)} 个）"}), 400

    # 随机打乱单词顺序
    random.shuffle(word_list)
    word_order = [w["id"] for w in word_list]

    cur = db.execute(
        """INSERT INTO quiz_sessions (total_count, word_order, current_index, status, mode)
           VALUES (?, ?, 0, 'in_progress', ?)""",
        (len(word_list), json.dumps(word_order), mode),
    )
    db.commit()
    session_id = cur.lastrowid

    # 返回第一道题
    question = _build_question(db, word_list, word_order, 0, session_id)
    return jsonify({
        "session_id": session_id,
        "total": len(word_list),
        "question": question,
    })


@app.route("/api/quiz/<int:session_id>/question", methods=["GET"])
def quiz_question(session_id):
    """获取当前题目（用于刷新/恢复）"""
    db = get_db()
    session = db.execute(
        "SELECT * FROM quiz_sessions WHERE id = ?", (session_id,)
    ).fetchone()

    if not session:
        return jsonify({"error": "会话不存在"}), 404
    if session["status"] != "in_progress":
        return jsonify({"error": "筛查已结束", "finished": True}), 200

    word_order = json.loads(session["word_order"])
    idx = session["current_index"]
    word_ids = [db.execute("SELECT id, english, chinese FROM words WHERE id = ?", (wid,)).fetchone() for wid in word_order]
    word_list = [dict(r) for r in word_ids if r]

    if idx >= len(word_list):
        return jsonify({"finished": True, "session_id": session_id})

    question = _build_question(db, word_list, [w["id"] for w in word_list], idx, session_id)
    return jsonify({
        "session_id": session_id,
        "current": idx + 1,
        "total": len(word_list),
        "question": question,
    })


@app.route("/api/quiz/<int:session_id>/answer", methods=["POST"])
def quiz_answer(session_id):
    """提交答案"""
    db = get_db()
    session = db.execute(
        "SELECT * FROM quiz_sessions WHERE id = ?", (session_id,)
    ).fetchone()

    if not session:
        return jsonify({"error": "会话不存在"}), 404
    if session["status"] != "in_progress":
        return jsonify({"error": "筛查已结束"}), 400

    data = request.get_json(silent=True) or {}
    word_id = data.get("word_id")
    selected = (data.get("selected") or "").strip()

    if not word_id or not selected:
        return jsonify({"error": "缺少 word_id 或 selected"}), 400

    # 判断对错
    word = db.execute("SELECT * FROM words WHERE id = ?", (word_id,)).fetchone()
    if not word:
        return jsonify({"error": "单词不存在"}), 404

    is_correct = word["chinese"].strip() == selected

    # 记录答案
    db.execute(
        "INSERT INTO quiz_answers (session_id, word_id, is_correct, user_answer) VALUES (?, ?, ?, ?)",
        (session_id, word_id, is_correct, selected),
    )

    # 更新会话计数
    if is_correct:
        db.execute("UPDATE quiz_sessions SET correct_count = correct_count + 1 WHERE id = ?", (session_id,))
    else:
        db.execute("UPDATE quiz_sessions SET wrong_count = wrong_count + 1 WHERE id = ?", (session_id,))

    # 推进索引
    word_order = json.loads(session["word_order"])
    new_index = session["current_index"] + 1
    finished = new_index >= len(word_order)

    if finished:
        db.execute(
            "UPDATE quiz_sessions SET current_index = ?, status = 'completed', completed_at = ? WHERE id = ?",
            (new_index, datetime.now(), session_id),
        )
    else:
        db.execute(
            "UPDATE quiz_sessions SET current_index = ? WHERE id = ?",
            (new_index, session_id),
        )

    db.commit()

    # 生成下一题
    response = {
        "is_correct": is_correct,
        "correct_answer": word["chinese"],
        "current": new_index,
        "total": session["total_count"],
    }

    if finished:
        response["finished"] = True
        # 最终统计
        final = db.execute("SELECT * FROM quiz_sessions WHERE id = ?", (session_id,)).fetchone()
        response["result"] = {
            "total": final["total_count"],
            "correct": final["correct_count"],
            "wrong": final["wrong_count"],
            "accuracy": round(final["correct_count"] / final["total_count"] * 100, 1) if final["total_count"] > 0 else 0,
        }
    else:
        word_list_ids = [db.execute("SELECT id, english, chinese FROM words WHERE id = ?", (wid,)).fetchone() for wid in word_order]
        word_list = [dict(r) for r in word_list_ids if r]
        response["next_question"] = _build_question(
            db, word_list, [w["id"] for w in word_list], new_index, session_id
        )

    return jsonify(response)


@app.route("/api/quiz/<int:session_id>/result", methods=["GET"])
def quiz_result(session_id):
    """获取筛查结果"""
    db = get_db()
    session = db.execute(
        "SELECT * FROM quiz_sessions WHERE id = ?", (session_id,)
    ).fetchone()

    if not session:
        return jsonify({"error": "会话不存在"}), 404

    answers = db.execute(
        """SELECT qa.*, w.english, w.chinese
           FROM quiz_answers qa
           JOIN words w ON qa.word_id = w.id
           WHERE qa.session_id = ?
           ORDER BY qa.created_at""",
        (session_id,),
    ).fetchall()

    return jsonify({
        "session": dict(session),
        "answers": [dict(a) for a in answers],
    })


def _build_question(db, word_list, word_order, index, session_id):
    """为第 index 个单词构建四选一题目"""
    current_word_id = word_order[index]
    current_word = next((w for w in word_list if w["id"] == current_word_id), None)
    if not current_word:
        return None

    correct_answer = current_word["chinese"]

    # 从其他单词中选 3 个干扰项
    other_chinese = [
        w["chinese"] for w in word_list if w["id"] != current_word_id
    ]
    # 去重（同一个中文释义可能对应多个英文单词）
    other_chinese = list(set(other_chinese))

    if len(other_chinese) >= 3:
        distractors = random.sample(other_chinese, 3)
    else:
        distractors = other_chinese
        # 不足 3 个时用填充
        fills = ["——", "……", "——"]
        for f in fills:
            if len(distractors) < 3 and f not in distractors and f != correct_answer:
                distractors.append(f)

    options = [correct_answer] + distractors
    random.shuffle(options)

    return {
        "word_id": current_word_id,
        "english": current_word["english"],
        "options": options,
    }


# ═══════════════════════════════════════════════════════════════
#  错题本 API
# ═══════════════════════════════════════════════════════════════

@app.route("/api/errors", methods=["GET"])
def get_errors():
    """获取错题本 — 所有答错过的单词，按错误次数降序"""
    db = get_db()
    rows = db.execute("""
        SELECT w.id, w.english, w.chinese, COUNT(qa.id) AS error_count
        FROM words w
        INNER JOIN quiz_answers qa ON w.id = qa.word_id AND qa.is_correct = 0
        GROUP BY w.id
        ORDER BY error_count DESC
    """).fetchall()

    return jsonify([dict(r) for r in rows])


# ═══════════════════════════════════════════════════════════════
#  导出 API
# ═══════════════════════════════════════════════════════════════

def _get_export_data(export_type):
    """获取导出数据"""
    db = get_db()
    if export_type == "errors":
        rows = db.execute("""
            SELECT w.english, w.chinese, COUNT(qa.id) AS error_count,
                   CASE WHEN COUNT(qa.id) > 0 THEN '未掌握' ELSE '已掌握' END AS status
            FROM words w
            INNER JOIN quiz_answers qa ON w.id = qa.word_id AND qa.is_correct = 0
            GROUP BY w.id
            ORDER BY error_count DESC
        """).fetchall()
    else:
        rows = db.execute("""
            SELECT w.english, w.chinese,
                   COALESCE(e.error_count, 0) AS error_count,
                   CASE WHEN e.error_count > 0 THEN '未掌握' ELSE '未筛查' END AS status
            FROM words w
            LEFT JOIN (
                SELECT word_id, COUNT(*) AS error_count
                FROM quiz_answers
                WHERE is_correct = 0
                GROUP BY word_id
            ) e ON w.id = e.word_id
            ORDER BY e.error_count DESC, w.created_at DESC
        """).fetchall()

    return [dict(r) for r in rows]


@app.route("/export/pdf", methods=["GET"])
def export_pdf():
    """导出 PDF"""
    export_type = request.args.get("type", "errors")
    data = _get_export_data(export_type)

    if not data:
        return jsonify({"error": "没有可导出的数据"}), 400

    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # 尝试注册中文字体
    font_registered = False
    font_paths = [
        ("C:/Windows/Fonts/msyh.ttc", "Microsoft YaHei"),
        ("C:/Windows/Fonts/msyhbd.ttc", "Microsoft YaHei"),
        ("C:/Windows/Fonts/simsun.ttc", "SimSun"),
        ("C:/Windows/Fonts/simhei.ttf", "SimHei"),
        ("/System/Library/Fonts/PingFang.ttc", "PingFang"),
        ("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc", "NotoSansCJK"),
    ]

    chinese_font = "Helvetica"
    for path, name in font_paths:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(name, path))
                chinese_font = name
                font_registered = True
                break
            except Exception:
                continue

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
        leftMargin=12 * mm,
        rightMargin=12 * mm,
    )

    styles = getSampleStyleSheet()
    chinese_style = ParagraphStyle(
        "Chinese",
        parent=styles["Normal"],
        fontName=chinese_font,
        fontSize=10,
        leading=16,
    )

    title_text = "单词筛查错题本" if export_type == "errors" else "单词筛查词汇表"
    title = Paragraph(f"<b>{title_text}</b>", ParagraphStyle(
        "Title2",
        parent=styles["Title"],
        fontName=chinese_font,
        fontSize=18,
        spaceAfter=12,
    ))

    # 表格数据
    table_data = [["序号", "英文", "中文", "掌握状态", "错误次数"]]
    for i, row in enumerate(data, 1):
        table_data.append([
            str(i),
            row["english"],
            row["chinese"],
            row.get("status", ""),
            str(row.get("error_count", 0)),
        ])

    col_widths = [30, 120, 200, 60, 45]
    table = Table(table_data, colWidths=col_widths, repeatRows=1)
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), chinese_font),
        ("FONTSIZE", (0, 0), (-1, 0), 11),
        ("FONTSIZE", (0, 1), (-1, -1), 10),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4A90D9")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
        ("ALIGN", (3, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F8FC")]),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))

    elements = [title, Spacer(1, 6 * mm), table]
    doc.build(elements)

    buf.seek(0)
    filename = f"单词筛查_{'错题本' if export_type == 'errors' else '词汇表'}.pdf"
    return send_file(buf, mimetype="application/pdf", as_attachment=True, download_name=filename)


@app.route("/export/word", methods=["GET"])
def export_word():
    """导出 Word 文档"""
    export_type = request.args.get("type", "errors")
    data = _get_export_data(export_type)

    if not data:
        return jsonify({"error": "没有可导出的数据"}), 400

    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # 标题
    title_text = "单词筛查错题本" if export_type == "errors" else "单词筛查词汇表"
    title = doc.add_heading(title_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(18)

    # 副标题：日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"导出日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # 空行

    # 表格
    table = doc.add_table(rows=1, cols=5, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    headers = ["序号", "英文", "中文", "掌握状态", "错误次数"]
    header_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        header_cells[i].text = text
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)
        # 设置表头背景色
        from docx.oxml.ns import qn
        shading_elm = header_cells[i]._element.get_or_add_tcPr()
        shading = shading_elm.makeelement(qn("w:shd"), {
            qn("w:fill"): "4A90D9",
            qn("w:val"): "clear",
        })
        shading_elm.append(shading)

    # 数据行
    for i, row in enumerate(data, 1):
        row_cells = table.add_row().cells
        values = [
            str(i),
            row["english"],
            row["chinese"],
            row.get("status", ""),
            str(row.get("error_count", 0)),
        ]
        for j, val in enumerate(values):
            row_cells[j].text = val
            for paragraph in row_cells[j].paragraphs:
                if j in (0, 3, 4):  # 序号、状态、次数居中
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # 设置列宽
    widths = [Cm(1.2), Cm(3.5), Cm(7), Cm(2.5), Cm(2)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"单词筛查_{'错题本' if export_type == 'errors' else '词汇表'}.docx"
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


# ═══════════════════════════════════════════════════════════════
#  离线缓存 Service Worker
# ═══════════════════════════════════════════════════════════════

@app.route("/sw.js")
def service_worker():
    """Service Worker — 缓存页面，支持 iPhone 离线使用"""
    return app.send_static_file("sw.js")


@app.route("/offline")
def offline_page():
    """离线版入口 — 预填词汇数据"""
    html_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "单词器_iPhone离线版.html")
    with open(html_path, "r", encoding="utf-8") as f:
        html = f.read()

    # 读取词汇
    db = get_db()
    rows = db.execute("SELECT english, chinese FROM words ORDER BY created_at DESC").fetchall()
    words_json = json.dumps([{"e": r["english"], "c": r["chinese"]} for r in rows], ensure_ascii=False)

    # 用 data script 嵌入到正确位置（最后 </body> 之前）
    injected = '<script id="__words_data__" type="application/json">' + words_json + '\n</script>'
    # 只替换最后一个 </body>
    idx = html.rfind("</body>")
    if idx >= 0:
        html = html[:idx] + injected + "\n" + html[idx:]

    from flask import Response
    resp = Response(html, mimetype="text/html")
    resp.headers["Cache-Control"] = "public, max-age=86400"
    return resp


# ─── 启动 ──────────────────────────────────────────────────
if __name__ == "__main__":
    print("=" * 50)
    print("  单词筛查背诵应用")
    print("=" * 50)
    print()
    print("  启动成功！")
    print("  电脑访问:  http://localhost:5000")
    print("  手机访问:  http://<本机IP>:5000")
    print("  按 Ctrl+C 停止")
    print("=" * 50)
    app.run(host="0.0.0.0", port=5000, debug=True)
